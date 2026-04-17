"""
FileProcessorTool for GazccThinking ReAct Agent
Handles: PDF, DOCX, CSV/XLSX, Images (OCR), Code, ZIP, Unknown/Binary
"""

import os
import io
import mimetypes
from pathlib import Path


# ─────────────────────────────────────────────
#  Constants
# ─────────────────────────────────────────────
MAX_FILE_SIZE     = 10 * 1024 * 1024   # 10 MB
SAMPLE_SIZE       = 2 * 1024 * 1024    # 2 MB – used when file exceeds limit
TEXT_FALLBACK_KB  = 5 * 1024           # 5 KB – plain-text fallback read
MAX_ROWS_DISPLAY  = 20

CODE_EXTENSIONS = {
    ".py": "python", ".js": "javascript", ".ts": "typescript",
    ".html": "html", ".css": "css", ".java": "java",
    ".c": "c", ".cpp": "cpp", ".cs": "csharp", ".go": "go",
    ".rs": "rust", ".rb": "ruby", ".php": "php",
    ".sh": "bash", ".yaml": "yaml", ".yml": "yaml",
    ".json": "json", ".xml": "xml", ".sql": "sql",
    ".kt": "kotlin", ".swift": "swift", ".r": "r",
    ".md": "markdown", ".toml": "toml", ".ini": "ini",
}

ENTRY_POINTS = {
    "main.py", "app.py", "run.py", "server.py",
    "index.html", "index.js", "index.ts",
    "manage.py", "setup.py", "Makefile",
    "docker-compose.yml", "docker-compose.yaml", "Dockerfile",
    "README.md", "readme.md",
}


# ─────────────────────────────────────────────
#  Helper
# ─────────────────────────────────────────────
def _fmt_size(n: int) -> str:
    if n < 1024:
        return f"{n} B"
    elif n < 1024 ** 2:
        return f"{n / 1024:.1f} KB"
    return f"{n / 1024 ** 2:.1f} MB"


def _header(name: str, kind: str, size: int) -> str:
    return f"📁 FILE: {name} | {kind} | {_fmt_size(size)}\n{'─' * 60}\n"


# ─────────────────────────────────────────────
#  FileProcessorTool
# ─────────────────────────────────────────────
class FileProcessorTool:
    name        = "file_processor"
    description = "Convert uploaded files to text for LLM analysis"

    # ── public entry point ──────────────────────────────────────────────
    def run(self, file_path: str, user_query: str = "") -> dict:
        """
        Returns:
            {
                'observation': str,            # Full text for LLM
                'status':      'success'|'error',
                'metadata':    dict
            }
        """
        path = Path(file_path)

        # ── existence check ─────────────────────────────────────────────
        if not path.exists():
            return self._err(
                f"File not found: {file_path}\n"
                "Check the path and try again.",
                {"file_path": file_path}
            )

        file_size = path.stat().st_size
        ext       = path.suffix.lower()

        # ── dispatch ────────────────────────────────────────────────────
        try:
            if ext == ".pdf":
                return self._pdf(path, file_size)
            elif ext == ".docx":
                return self._docx(path, file_size)
            elif ext in (".csv", ".xlsx", ".xls"):
                return self._spreadsheet(path, file_size, ext)
            elif ext in (".jpg", ".jpeg", ".png", ".bmp", ".tiff", ".webp"):
                return self._image(path, file_size, user_query)
            elif ext in CODE_EXTENSIONS:
                return self._code(path, file_size, ext)
            elif ext == ".zip":
                return self._zip(path, file_size)
            else:
                return self._fallback(path, file_size)

        except Exception as exc:
            return self._err(
                f"Unexpected error processing {path.name}: {exc}",
                {"file_path": file_path, "exception": type(exc).__name__}
            )

    # ── PDF ─────────────────────────────────────────────────────────────
    def _pdf(self, path: Path, size: int) -> dict:
        try:
            import PyPDF2
        except ImportError:
            return self._import_err("PyPDF2")

        lines   = [_header(path.name, "PDF", size)]
        meta    = {"pages": 0, "title": None, "author": None}
        oversized = size > MAX_FILE_SIZE

        try:
            with open(path, "rb") as fh:
                reader = PyPDF2.PdfReader(fh)
                info   = reader.metadata or {}
                meta["title"]  = info.get("/Title")
                meta["author"] = info.get("/Author")
                meta["pages"]  = len(reader.pages)

                if info.get("/Title"):
                    lines.append(f"Title  : {info['/Title']}")
                if info.get("/Author"):
                    lines.append(f"Author : {info['/Author']}")
                lines.append(f"Pages  : {meta['pages']}\n")

                max_pages = 10 if oversized else meta["pages"]
                for i, page in enumerate(reader.pages[:max_pages], 1):
                    try:
                        text = page.extract_text() or "[No extractable text]"
                    except Exception as e:
                        text = f"[Error extracting page: {e}]"
                    lines.append(f"=== Page {i} ===\n{text.strip()}\n")

                if oversized and meta["pages"] > max_pages:
                    lines.append(
                        f"[⚠ File >10 MB — showing first {max_pages} of "
                        f"{meta['pages']} pages]"
                    )

        except PyPDF2.errors.PdfReadError as e:
            lines.append(f"[⚠ Corrupted PDF – partial read. Error: {e}]")

        return self._ok("\n".join(lines), meta)

    # ── DOCX ────────────────────────────────────────────────────────────
    def _docx(self, path: Path, size: int) -> dict:
        try:
            from docx import Document
        except ImportError:
            return self._import_err("python-docx")

        lines = [_header(path.name, "Word Document", size)]
        meta  = {"paragraphs": 0, "tables": 0}

        try:
            doc = Document(str(path))

            for para in doc.paragraphs:
                if para.text.strip():
                    meta["paragraphs"] += 1
                    style = para.style.name or ""
                    if "Heading 1" in style:
                        lines.append(f"\n# {para.text}")
                    elif "Heading 2" in style:
                        lines.append(f"\n## {para.text}")
                    elif "Heading 3" in style:
                        lines.append(f"\n### {para.text}")
                    else:
                        lines.append(para.text)

            for idx, table in enumerate(doc.tables, 1):
                meta["tables"] += 1
                lines.append(f"\n--- Table {idx} ---")
                for row in table.rows:
                    cells = [c.text.strip().replace("\n", " ") for c in row.cells]
                    lines.append("| " + " | ".join(cells) + " |")

        except Exception as e:
            lines.append(f"[⚠ Partial read – error at: {e}]")

        return self._ok("\n".join(lines), meta)

    # ── Spreadsheet ─────────────────────────────────────────────────────
    def _spreadsheet(self, path: Path, size: int, ext: str) -> dict:
        try:
            import pandas as pd
        except ImportError:
            return self._import_err("pandas")

        kind  = "CSV" if ext == ".csv" else "Excel Spreadsheet"
        lines = [_header(path.name, kind, size)]
        meta  = {}

        try:
            if ext == ".csv":
                if size > MAX_FILE_SIZE:
                    df = pd.read_csv(path, nrows=MAX_ROWS_DISPLAY)
                    lines.append(f"[⚠ File >10 MB — showing first {MAX_ROWS_DISPLAY} rows]\n")
                else:
                    df = pd.read_csv(path)
            else:
                try:
                    import openpyxl  # noqa: F401
                except ImportError:
                    return self._import_err("openpyxl")
                df = pd.read_excel(path, nrows=MAX_ROWS_DISPLAY if size > MAX_FILE_SIZE else None)

            meta = {
                "rows":    int(df.shape[0]),
                "columns": int(df.shape[1]),
                "dtypes":  {c: str(t) for c, t in df.dtypes.items()},
            }

            lines.append(f"Shape   : {df.shape[0]} rows × {df.shape[1]} columns")
            lines.append(f"Columns : {list(df.columns)}\n")
            lines.append("=== First rows ===")
            lines.append(df.head(MAX_ROWS_DISPLAY).to_string(index=True))

            numeric = df.select_dtypes(include="number")
            if not numeric.empty:
                lines.append("\n=== Numeric stats ===")
                lines.append(numeric.describe().to_string())

        except Exception as e:
            lines.append(f"[⚠ Read error: {e}]")

        return self._ok("\n".join(lines), meta)

    # ── Image / OCR ─────────────────────────────────────────────────────
    def _image(self, path: Path, size: int, user_query: str) -> dict:
        try:
            from PIL import Image
        except ImportError:
            return self._import_err("Pillow")

        lines = [_header(path.name, "Image", size)]
        meta  = {}

        try:
            img = Image.open(str(path))
            meta = {
                "width":  img.width,
                "height": img.height,
                "mode":   img.mode,
                "format": img.format,
            }
            lines.append(
                f"Dimensions : {img.width} × {img.height} px\n"
                f"Mode       : {img.mode}\n"
                f"Format     : {img.format}\n"
            )

            q = (user_query or "").lower()
            if any(w in q for w in ("code", "error", "screenshot", "terminal", "log")):
                lines.append(
                    "[ℹ Likely a screenshot — OCR output may contain code / error messages]\n"
                )

            # OCR
            try:
                import pytesseract
                ocr_text = pytesseract.image_to_string(img).strip()
                if ocr_text:
                    lines.append("=== OCR Text ===")
                    lines.append(ocr_text)
                else:
                    lines.append("[OCR: No text detected]")
            except ImportError:
                lines.append(
                    "[pytesseract not installed — OCR unavailable]\n"
                    "Install: pip install pytesseract  (also needs Tesseract binary)"
                )
            except Exception as e:
                lines.append(f"[OCR error: {e}]")

        except Exception as e:
            lines.append(f"[⚠ Image read error: {e}]")

        return self._ok("\n".join(lines), meta)

    # ── Code ────────────────────────────────────────────────────────────
    def _code(self, path: Path, size: int, ext: str) -> dict:
        lang  = CODE_EXTENSIONS.get(ext, "text")
        lines = [_header(path.name, f"Code ({lang})", size)]
        meta  = {"language": lang, "lines": 0}

        try:
            max_bytes = SAMPLE_SIZE if size > MAX_FILE_SIZE else None
            with open(path, "r", encoding="utf-8", errors="replace") as fh:
                content = fh.read(max_bytes)

            line_count = content.count("\n") + 1
            meta["lines"] = line_count

            lines.append(f"Language : {lang}")
            lines.append(f"Lines    : {line_count}")
            if size > MAX_FILE_SIZE:
                lines.append(f"[⚠ File >10 MB — showing first {_fmt_size(SAMPLE_SIZE)}]\n")
            lines.append(f"\n```{lang}\n{content}\n```")

        except Exception as e:
            lines.append(f"[⚠ Read error: {e}]")

        return self._ok("\n".join(lines), meta)

    # ── ZIP ─────────────────────────────────────────────────────────────
    def _zip(self, path: Path, size: int) -> dict:
        import zipfile

        lines = [_header(path.name, "ZIP Archive", size)]
        meta  = {"file_count": 0, "entry_points": []}

        try:
            with zipfile.ZipFile(str(path), "r") as zf:
                members = zf.infolist()
                meta["file_count"] = len(members)

                lines.append(f"Files : {len(members)}\n")
                lines.append("=== Contents ===")

                entry_points = []
                for info in members:
                    base = Path(info.filename).name
                    flag = " ← ENTRY POINT" if base in ENTRY_POINTS else ""
                    size_str = _fmt_size(info.file_size)
                    lines.append(f"  {info.filename:<50} {size_str}{flag}")
                    if flag:
                        entry_points.append(info.filename)

                meta["entry_points"] = entry_points

                if entry_points:
                    lines.append("\n=== Detected Entry Points ===")
                    for ep in entry_points:
                        lines.append(f"  • {ep}")

                        # Peek at small entry-point files
                        info = zf.getinfo(ep)
                        if info.file_size < 8192:
                            try:
                                content = zf.read(ep).decode("utf-8", errors="replace")
                                ext     = Path(ep).suffix.lower()
                                lang    = CODE_EXTENSIONS.get(ext, "text")
                                lines.append(f"\n--- {ep} ---")
                                lines.append(f"```{lang}\n{content}\n```")
                            except Exception:
                                pass

        except zipfile.BadZipFile:
            lines.append("[⚠ Corrupted ZIP – cannot read archive]")

        return self._ok("\n".join(lines), meta)

    # ── Fallback / Unknown ───────────────────────────────────────────────
    def _fallback(self, path: Path, size: int) -> dict:
        mime, _ = mimetypes.guess_type(str(path))
        kind    = mime or "Unknown/Binary"
        lines   = [_header(path.name, kind, size)]
        meta    = {"mime": mime}

        # Attempt UTF-8 text read
        try:
            with open(path, "r", encoding="utf-8", errors="strict") as fh:
                content = fh.read(TEXT_FALLBACK_KB)
            lines.append("=== Text Content (first 5 KB) ===")
            lines.append(content)
            meta["readable"] = True
            return self._ok("\n".join(lines), meta)
        except (UnicodeDecodeError, ValueError):
            pass

        # Hex preview + printable strings
        try:
            with open(path, "rb") as fh:
                raw = fh.read(512)

            hex_preview = "\n".join(
                " ".join(f"{b:02x}" for b in raw[i:i+16])
                for i in range(0, len(raw), 16)
            )
            printable = "".join(
                chr(b) if 32 <= b < 127 else "." for b in raw
            )

            # Extract "strings" (runs of 4+ printable chars)
            import re
            strings_found = re.findall(r"[ -~]{4,}", raw.decode("latin-1", errors="replace"))

            lines.append("=== Hex Preview (first 512 bytes) ===")
            lines.append(hex_preview)
            lines.append("\n=== ASCII Representation ===")
            lines.append(printable)
            if strings_found:
                lines.append("\n=== Extracted Strings ===")
                lines.append("\n".join(strings_found[:50]))

            meta["readable"] = False

        except Exception as e:
            lines.append(f"[⚠ Cannot read file: {e}]")

        return self._ok("\n".join(lines), meta)

    # ── response builders ────────────────────────────────────────────────
    @staticmethod
    def _ok(observation: str, metadata: dict) -> dict:
        return {"observation": observation, "status": "success", "metadata": metadata}

    @staticmethod
    def _err(message: str, metadata: dict) -> dict:
        return {"observation": f"❌ ERROR: {message}", "status": "error", "metadata": metadata}

    @staticmethod
    def _import_err(package: str) -> dict:
        return {
            "observation": (
                f"❌ Missing dependency: {package}\n"
                f"Install with: pip install {package}"
            ),
            "status": "error",
            "metadata": {"missing_package": package},
        }
